import os
import sys
try:
    import langchain
    print(f"DEBUG: LangChain version: {langchain.__version__}")
    print(f"DEBUG: LangChain path: {langchain.__file__}")
except ImportError:
    print("DEBUG: LangChain not found")

from langchain_huggingface import HuggingFaceEmbeddings
from langchain_ollama import ChatOllama
from langchain_classic.chains import RetrievalQA
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_chroma import Chroma
from hwp_loader import get_hwp_text
from langchain_core.documents import Document
from langchain_core.prompts import PromptTemplate
import time

class RagEngine:
    def __init__(self, persist_directory="./chroma_db"):
        self.persist_directory = persist_directory
        # Using a lightweight Korean embedding model
        self.embedding_model = HuggingFaceEmbeddings(
            model_name="jhgan/ko-sroberta-multitask"
        )
        # LLM 설정 (Ollama 로컬 모델)
        self.llm = ChatOllama(
            # model: 사용할 LLM 모델 지정
            # - exaone3.5:2.4b: 한국어 특화 2.4B 파라미터 모델 (LG AI 개발, GPU 4GB 가능)
            # - qwen2.5:7b-instruct-q4_K_M: 다국어 7B 모델 (더 강력하지만 메모리 더 필요)
            model="exaone3.5:2.4b",

            # temperature: 답변의 창의성/무작위성 조절 (0.0~1.0)
            # - 0.0에 가까울수록: 결정적이고 정확한 답변 (환각 감소, 문서 기반 Q&A에 적합)
            # - 1.0에 가까울수록: 창의적이고 다양한 답변 (창작 작업에 적합)
            temperature=0.1,

            # num_predict: 생성할 최대 토큰(단어) 수
            # - 512: 중간 길이 답변 (~300-400자 정도)
            # - 높을수록 긴 답변 가능하지만 처리 시간 증가
            num_predict=512,

            # top_p: 누적 확률 기반 토큰 샘플링 (0.0~1.0)
            # - 0.9: 상위 90% 확률의 토큰만 고려 (적절한 다양성 유지)
            # - 낮을수록 보수적, 높을수록 다양한 표현
            top_p=0.9,

            # repeat_penalty: 반복 방지 페널티 (1.0 이상)
            # - 1.0: 페널티 없음
            # - 1.1: 약간의 반복 방지 (같은 단어/문장 반복 감소)
            # - 너무 높으면 부자연스러운 답변 생성 가능
            repeat_penalty=1.1
        ) 
        self.vectorstore = None


    def load_documents(self, doc_paths):
        """
        doc_paths: List of file or directory paths
        """
        documents = []
        if isinstance(doc_paths, str):
            doc_paths = [doc_paths]

        for path in doc_paths:
            if os.path.isdir(path):
                for root, dirs, files in os.walk(path):
                    for file in files:
                        file_path = os.path.join(root, file)
                        self._load_single_file(file_path, documents)
            elif os.path.isfile(path):
                self._load_single_file(path, documents)
        
        return documents

    def _load_single_file(self, file_path, documents):
        try:
            text = ""
            if file_path.lower().endswith(".docx"):
                # DOCX 파일 (최우선)
                from docx import Document as DocxDocument
                doc = DocxDocument(file_path)

                text_parts = []

                # 문단 읽기 (세로쓰기 처리)
                buffer = []
                for para in doc.paragraphs:
                    para_text = para.text.strip()
                    if para_text:
                        # 한 글자만 있는 경우 버퍼에 모으기
                        if len(para_text) <= 2:
                            buffer.append(para_text)
                        else:
                            # 버퍼에 모인 내용 먼저 추가
                            if buffer:
                                text_parts.append(''.join(buffer))
                                buffer = []
                            text_parts.append(para_text)

                # 남은 버퍼 처리
                if buffer:
                    text_parts.append(''.join(buffer))

                # 표(table) 읽기
                for table in doc.tables:
                    for row in table.rows:
                        row_text = []
                        for cell in row.cells:
                            cell_text = cell.text.strip()
                            if cell_text:
                                row_text.append(cell_text)
                        if row_text:
                            text_parts.append(' | '.join(row_text))

                text = '\n'.join(text_parts)
            elif file_path.lower().endswith(".txt"):
                # TXT 파일
                with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                    text = f.read()
            elif file_path.lower().endswith(".hwp"):
                # HWP 파일 (문제 있음)
                text = get_hwp_text(file_path)
            # MD 파일은 제외
            # elif file_path.lower().endswith(".md"):
            #     with open(file_path, "r", encoding="utf-8") as f:
            #         text = f.read()

            if text.strip():
                documents.append(Document(page_content=text, metadata={"source": os.path.basename(file_path)}))
                print(f"Loaded: {file_path}")
        except Exception as e:
            print(f"Error loading {file_path}: {e}")

    def create_index(self, documents):
        if not documents:
            print("No documents to index.")
            return

        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=2000,  # 청크 크기 증가 (제목 + 여러 조문 포함)
            chunk_overlap=300,  # 중복 영역 확대 (제목이 다음 청크에도 포함되도록)
            separators=["\n\n", "\n", ".", " ", ""]  # 자연스러운 구분점에서 분할
        )
        texts = text_splitter.split_documents(documents)
        
        self.vectorstore = Chroma.from_documents(
            documents=texts,
            embedding=self.embedding_model,
            persist_directory=self.persist_directory
        )
        # Chroma automatically persists in newer versions, but explicit persist doesn't hurt if old version
        # self.vectorstore.persist() 
        print(f"Indexed {len(texts)} chunks.")

    def load_index(self):
        if os.path.exists(self.persist_directory):
            self.vectorstore = Chroma(
                persist_directory=self.persist_directory,
                embedding_function=self.embedding_model
            )
            return True
        return False

    def ask(self, query):
        if not self.vectorstore:
            # Try to load if not loaded
            if not self.load_index():
                return {"result": "문서가 인덱싱되지 않았습니다. 먼저 문서를 로드해주세요.", "source_documents": []}

        print(f"[LOG] 질의: {query}")
        start = time.time()

        print(f"[LOG] 벡터 검색 시작...")
        search_start = time.time()

        retriever = self.vectorstore.as_retriever(
            search_type="similarity",  # 유사도 검색으로 변경 (관련성 우선)
            search_kwargs={
                "k": 5  # 가장 관련 높은 5개로 증가
            }
        )
        docs = retriever.invoke(query)

        print(f"[LOG] 벡터 검색 완료 ({time.time() - search_start:.2f}초)")
        print(f"[LOG] 검색된 문서 수: {len(docs)}개")
        print("=" * 80)
        for i, doc in enumerate(docs, 1):
            print(f"\n[LOG] 문서 {i}: {doc.metadata.get('source', 'Unknown')} (길이: {len(doc.page_content)} 글자)")
            # 검색된 내용 미리보기 (앞 500자)
            preview = doc.page_content[:500].replace('\n', ' ')
            print(f"[LOG] 내용 미리보기: {preview}...")
        print("=" * 80)

        prompt_template = """아래 문서 내용을 읽고 질문에 답하세요.

문서:
{context}

질문: {question}

규칙:
- 문서에 있는 내용만 답변하세요
- 문서에 "제65조"라고 쓰여있으면 그대로 "제65조"라고 답변하세요 (번호를 절대 바꾸지 마세요)
- 질문에 "2장"이 있으면 문서에서 "제2장" 또는 "제 2 장"을 찾으세요
- 문서에 없으면 "문서에 해당 내용이 없습니다"라고 답변하세요
- 추측하지 마세요

답변:"""

        PROMPT = PromptTemplate(
            template=prompt_template, input_variables=["context", "question"]
        )

        qa_chain = RetrievalQA.from_chain_type(
            llm=self.llm,
            chain_type="stuff",
            retriever=retriever,
            return_source_documents=True,
            chain_type_kwargs={"prompt": PROMPT}
        )

        print(f"[LOG] LLM 응답 생성 시작...")
        llm_start = time.time()

        result = qa_chain.invoke({"query": query})

        print(f"[LOG] LLM 응답 완료 ({time.time() - llm_start:.2f}초)")
        print("=" * 80)
        print(f"[LOG] LLM 답변:\n{result.get('result', '답변 없음')}")
        print("=" * 80)
        print(f"[LOG] 전체 소요 시간: {time.time() - start:.2f}초")

        return result
