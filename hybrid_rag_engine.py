"""
하이브리드 검색 RAG 엔진 (수동 구현)
BM25 (키워드 60%) + Vector (의미 40%) 검색 결합
"""
import os
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_ollama import ChatOllama
from langchain_classic.chains import RetrievalQA
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_chroma import Chroma
from hwp_loader import get_hwp_text
from langchain_core.documents import Document
from langchain_core.prompts import PromptTemplate
from langchain_community.retrievers import BM25Retriever
import time

class HybridRagEngine:
    def __init__(self, persist_directory="./chroma_db"):
        self.persist_directory = persist_directory
        self.embedding_model = HuggingFaceEmbeddings(
            model_name="jhgan/ko-sroberta-multitask"
        )
        self.llm = ChatOllama(
            model="qwen2.5:3b",
            temperature=0.7,
            num_predict=512,
            top_p=0.9,
            repeat_penalty=1.1
        ) 
        self.vectorstore = None
        self.bm25_retriever = None
        self.all_splits = []

    def load_documents(self, doc_paths):
        """문서 로드"""
        documents = []
        if isinstance(doc_paths, str):
            doc_paths = [doc_paths]

        for path in doc_paths:
            if os.path.isdir(path):
                for root, dirs, files in os.walk(path):
                    for file in files:
                        file_path = os.path.join(root, file)
                        self._load_single_file(file_path, documents)
            elif os.path.isfile(path):
                self._load_single_file(path, documents)
        
        return documents

    def _load_single_file(self, file_path, documents):
        """단일 파일 로드"""
        try:
            text = ""
            if file_path.lower().endswith(".docx"):
                from docx import Document as DocxDocument
                doc = DocxDocument(file_path)
                text_parts = []
                buffer = []
                for para in doc.paragraphs:
                    para_text = para.text.strip()
                    if para_text:
                        if len(para_text) <= 2:
                            buffer.append(para_text)
                        else:
                            if buffer:
                                text_parts.append(''.join(buffer))
                                buffer = []
                            text_parts.append(para_text)
                if buffer:
                    text_parts.append(''.join(buffer))
                for table in doc.tables:
                    for row in table.rows:
                        row_text = []
                        for cell in row.cells:
                            cell_text = cell.text.strip()
                            if cell_text:
                                row_text.append(cell_text)
                        if row_text:
                            text_parts.append(' | '.join(row_text))
                text = '\n'.join(text_parts)
            elif file_path.lower().endswith(".txt"):
                with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                    text = f.read()
            elif file_path.lower().endswith(".hwp"):
                text = get_hwp_text(file_path)

            if text.strip():
                documents.append(Document(page_content=text, metadata={"source": os.path.basename(file_path)}))
                print(f"Loaded: {file_path}")
        except Exception as e:
            print(f"Error loading {file_path}: {e}")

    def create_index(self, documents):
        """하이브리드 인덱스 생성"""
        if not documents:
            print("No documents to index.")
            return

        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=2000,
            chunk_overlap=300,
            separators=["\n\n", "\n", ".", " ", ""]
        )
        self.all_splits = text_splitter.split_documents(documents)
        
        # 벡터 인덱스
        self.vectorstore = Chroma.from_documents(
            documents=self.all_splits,
            embedding=self.embedding_model,
            persist_directory=self.persist_directory
        )
        
        # BM25 인덱스
        self.bm25_retriever = BM25Retriever.from_documents(self.all_splits)
        self.bm25_retriever.k = 5
        
        print(f"✅ Indexed {len(self.all_splits)} chunks (Hybrid: BM25 + Vector)")

    def load_index(self):
        """기존 인덱스 로드"""
        if os.path.exists(self.persist_directory):
            self.vectorstore = Chroma(
                persist_directory=self.persist_directory,
                embedding_function=self.embedding_model
            )
            return True
        return False

    def _hybrid_search(self, query, k=5):
        """하이브리드 검색: BM25 + Vector 결합"""
        # BM25 검색
        if self.bm25_retriever and self.all_splits:
            self.bm25_retriever = BM25Retriever.from_documents(self.all_splits)
            self.bm25_retriever.k = k
            bm25_docs = self.bm25_retriever.invoke(query)
        else:
            bm25_docs = []
        
        # 벡터 검색
        vector_docs = self.vectorstore.as_retriever(
            search_kwargs={"k": k}
        ).invoke(query)
        
        # 결과 합치기 (가중치: BM25 60%, Vector 40%)
        # 중복 제거하면서 순위 조정
        seen_contents = set()
        merged_docs = []
        
        # BM25 결과 먼저 추가 (높은 가중치)
        for doc in bm25_docs[:3]:  # 상위 3개
            content_hash = hash(doc.page_content)
            if content_hash not in seen_contents:
                seen_contents.add(content_hash)
                merged_docs.append(doc)
        
        # Vector 결과 추가 (낮은 가중치)
        for doc in vector_docs[:3]:  # 상위 3개
            content_hash = hash(doc.page_content)
            if content_hash not in seen_contents:
                seen_contents.add(content_hash)
                merged_docs.append(doc)
        
        # 부족하면 나머지 추가
        for doc in bm25_docs[3:] + vector_docs[3:]:
            if len(merged_docs) >= k:
                break
            content_hash = hash(doc.page_content)
            if content_hash not in seen_contents:
                seen_contents.add(content_hash)
                merged_docs.append(doc)
        
        return merged_docs[:k]

    def ask(self, query):
        """하이브리드 검색으로 질의응답"""
        if not self.vectorstore:
            if not self.load_index():
                return {"result": "문서가 인덱싱되지 않았습니다.", "source_documents": []}

        print(f"[LOG] 질의: {query}")
        start = time.time()

        # 하이브리드 검색
        if self.bm25_retriever and self.all_splits:
            print("[LOG] 하이브리드 검색 (BM25 60% + Vector 40%)")
            docs = self._hybrid_search(query, k=5)
        else:
            print("[LOG] 벡터 검색만 사용 (BM25 인덱스 없음)")
            docs = self.vectorstore.as_retriever(
                search_kwargs={"k": 5}
            ).invoke(query)
        
        print(f"[LOG] 검색 완료 ({time.time() - start:.2f}초)")
        print(f"[LOG] 검색된 문서 수: {len(docs)}개")
        
        for i, doc in enumerate(docs, 1):
            print(f"[LOG] 문서 {i}: {doc.metadata.get('source', 'Unknown')} (길이: {len(doc.page_content)} 글자)")

        # 커스텀 리트리버 래퍼
        class CustomRetriever:
            def __init__(self, docs):
                self.docs = docs
            
            def invoke(self, query):
                return self.docs
            
            def get_relevant_documents(self, query):
                return self.docs

        retriever = CustomRetriever(docs)

        prompt_template = """당신은 한국의료연구원의 사내 규정 전문가입니다.

[중요] 질문에서 특정 장/조 번호를 요청했는데 문서에 없다면:
"제공된 문서에 제X장(제X조)에 대한 내용이 없습니다."라고 답변하세요.

[문서 내용]
{context}

[질문]
{question}

[답변 규칙]
1. 요청한 장/조 번호가 문서에 정확히 있는지 확인
2. 없으면 "문서에 없습니다" 명확히 답변
3. 있으면 번호와 제목을 먼저 명시하고 내용 요약
4. 절대 다른 조문으로 대체하지 말 것

답변:"""

        PROMPT = PromptTemplate(
            template=prompt_template, input_variables=["context", "question"]
        )

        qa_chain = RetrievalQA.from_chain_type(
            llm=self.llm,
            chain_type="stuff",
            retriever=retriever,
            return_source_documents=True,
            chain_type_kwargs={"prompt": PROMPT}
        )

        print(f"[LOG] LLM 응답 생성 시작...")
        llm_start = time.time()

        result = qa_chain.invoke({"query": query})

        print(f"[LOG] LLM 응답 완료 ({time.time() - llm_start:.2f}초)")
        print(f"[LOG] 전체 소요 시간: {time.time() - start:.2f}초")

        return result
