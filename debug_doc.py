from docx import Document
import os

doc_path = r"D:\work\group\local_doc_search\doc\규정(2025.11.04.) (2).docx"

print("=" * 80)
print("문서 내용 확인")
print("=" * 80)

doc = Document(doc_path)

all_text = []
for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    if text:
        all_text.append(text)
        if "15조" in text or "제15조" in text:
            print(f"\n[발견!] 문단 {i}: {text[:100]}")

# 표에서도 확인
for table_idx, table in enumerate(doc.tables):
    for row_idx, row in enumerate(table.rows):
        for cell_idx, cell in row.cells:
            text = cell.text.strip()
            if "15조" in text or "제15조" in text:
                print(f"\n[발견!] 표 {table_idx}, 행 {row_idx}, 셀 {cell_idx}: {text[:100]}")

# 전체 텍스트에서 15조 찾기
full_text = '\n'.join(all_text)
if "15조" in full_text or "제15조" in full_text:
    print("\n✓ 문서에 '15조' 또는 '제15조'가 포함되어 있습니다.")

    # 15조 주변 내용 출력
    lines = full_text.split('\n')
    for i, line in enumerate(lines):
        if "15조" in line or "제15조" in line:
            start = max(0, i-2)
            end = min(len(lines), i+5)
            print(f"\n[주변 내용 (라인 {i})]:")
            print('\n'.join(lines[start:end]))
            print("-" * 80)
else:
    print("\n✗ 문서에 '15조' 또는 '제15조'가 없습니다.")

print(f"\n총 문단 수: {len(all_text)}")
print(f"총 문자 수: {len(full_text)}")
